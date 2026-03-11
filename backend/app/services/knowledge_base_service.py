"""
Knowledge Base Service — Document ingestion, chunking, embedding, and RAG retrieval

Manages the full RAG pipeline:
  1. Document upload and parsing (PDF, DOCX, TXT)
  2. Recursive text chunking with configurable overlap
  3. Vector embedding via local or remote embedding service
  4. Similarity search using pgvector cosine distance

Uses pgvector on existing PostgreSQL — no additional infrastructure needed.
"""

import asyncio
import io
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional

from sqlalchemy import func, select, delete, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.knowledge_base import KBDocument, KBChunk
from app.services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ChunkResult:
    """A single search result from the knowledge base."""
    chunk_id: int
    document_id: int
    document_title: str
    content: str
    chunk_index: int
    page_number: Optional[int]
    score: float  # Cosine similarity (0-1, higher is better)
    metadata: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "document_title": self.document_title,
            "content": self.content,
            "chunk_index": self.chunk_index,
            "page_number": self.page_number,
            "score": round(self.score, 4),
            "metadata": self.metadata,
        }


@dataclass
class DocumentInfo:
    """Summary information about an uploaded document."""
    id: int
    title: str
    filename: str
    file_type: str
    file_size: Optional[int]
    status: str
    chunk_count: int
    category: Optional[str]
    created_at: Optional[datetime]

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "title": self.title,
            "filename": self.filename,
            "file_type": self.file_type,
            "file_size": self.file_size,
            "status": self.status,
            "chunk_count": self.chunk_count,
            "category": self.category,
            "created_at": self.created_at.isoformat() if self.created_at else None,
        }


# ============================================================================
# Text Chunking
# ============================================================================

def chunk_text(
    text_content: str,
    chunk_size: int = 1024,
    chunk_overlap: int = 200,
) -> List[Dict[str, Any]]:
    """Split text into overlapping chunks using recursive character splitting.

    Tries to split on paragraph boundaries first, then sentences, then words.

    Args:
        text_content: The full text to chunk.
        chunk_size: Maximum characters per chunk.
        chunk_overlap: Number of overlapping characters between chunks.

    Returns:
        List of dicts with 'content', 'start_char', 'end_char' keys.
    """
    if not text_content or not text_content.strip():
        return []

    separators = ["\n\n", "\n", ". ", " "]
    chunks = []
    start = 0

    while start < len(text_content):
        end = start + chunk_size

        if end >= len(text_content):
            # Last chunk — take everything remaining
            chunk = text_content[start:].strip()
            if chunk:
                chunks.append({
                    "content": chunk,
                    "start_char": start,
                    "end_char": len(text_content),
                })
            break

        # Try to find a good split point using separators
        split_pos = end
        for sep in separators:
            # Look backwards from end for the separator
            pos = text_content.rfind(sep, start + chunk_overlap, end)
            if pos > start:
                split_pos = pos + len(sep)
                break

        chunk = text_content[start:split_pos].strip()
        if chunk:
            chunks.append({
                "content": chunk,
                "start_char": start,
                "end_char": split_pos,
            })

        # Move forward with overlap
        start = split_pos - chunk_overlap
        if start <= chunks[-1]["start_char"] if chunks else 0:
            # Prevent infinite loop on pathological input
            start = split_pos

    return chunks


# ============================================================================
# Document Parsing
# ============================================================================

def parse_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from PDF bytes. Falls back to OCR for scanned PDFs.

    Returns (text, page_count).
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        text = "\n\n".join(pages)
        page_count = len(reader.pages)

        # If PyPDF2 extracted meaningful text, return it
        if text.strip() and len(text.strip().split()) > 50:
            return text, page_count

        # Scanned PDF — fall back to OCR
        logger.info(f"PDF has little/no extractable text ({len(text.split())} words), trying OCR...")
        return _ocr_pdf(file_bytes, page_count)

    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise ValueError(f"Failed to parse PDF: {e}")


def _ocr_pdf(file_bytes: bytes, page_count_hint: int = 0) -> tuple[str, int]:
    """OCR a scanned PDF using Tesseract + pdf2image."""
    try:
        import tempfile
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as ie:
        logger.warning(f"OCR dependencies not available: {ie}")
        return "", page_count_hint

    try:
        images = convert_from_bytes(file_bytes, dpi=150, fmt="png")
        all_text = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang="eng").strip()
            if text:
                all_text.append(text)
        full_text = "\n\n".join(all_text)
        full_text = full_text.replace("\x00", "")
        logger.info(f"OCR extracted {len(full_text.split())} words from {len(images)} pages")
        return full_text, len(images)
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return "", page_count_hint


def parse_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from DOCX bytes. Returns (text, paragraph_count)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs), len(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Failed to parse DOCX: {e}")


def parse_text(file_bytes: bytes) -> tuple[str, int]:
    """Decode plain text/markdown bytes. Returns (text, line_count)."""
    content = file_bytes.decode("utf-8", errors="replace")
    lines = content.split("\n")
    return content, len(lines)


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "txt": parse_text,
    "md": parse_text,
    "csv": parse_text,
}


# ============================================================================
# HTML Helpers (for URL ingestion)
# ============================================================================

def _extract_html_text(html: str) -> str:
    """Extract readable text from HTML. Uses BeautifulSoup if available."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        pass
    # Fallback: strip HTML tags with regex
    import re
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_html_title(html: str) -> Optional[str]:
    """Extract <title> tag content from HTML."""
    import re
    match = re.search(r"<title[^>]*>(.*?)</title>", html, re.DOTALL | re.IGNORECASE)
    if match:
        return re.sub(r"\s+", " ", match.group(1)).strip()[:200]
    return None


# ============================================================================
# Knowledge Base Service
# ============================================================================

class KnowledgeBaseService:
    """Manages document ingestion, embedding, and RAG retrieval.

    Async-first design. All database operations go through the injected session.
    """

    def __init__(self, db: AsyncSession, tenant_id: int):
        self.db = db
        self.tenant_id = tenant_id
        self._embedding_service: Optional[EmbeddingService] = None

    @property
    def embedding_service(self) -> EmbeddingService:
        if self._embedding_service is None:
            self._embedding_service = EmbeddingService()
        return self._embedding_service

    # ------------------------------------------------------------------
    # Document Ingestion
    # ------------------------------------------------------------------

    async def ingest_document(
        self,
        file_bytes: bytes,
        filename: str,
        title: Optional[str] = None,
        category: Optional[str] = None,
        description: Optional[str] = None,
        tags: Optional[List[str]] = None,
        uploaded_by: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Parse, chunk, embed, and store a document.

        Args:
            file_bytes: Raw file content.
            filename: Original filename (used to detect file type).
            title: Human-readable title (defaults to filename).
            category: Document category (annual_report, quarterly_report, etc.).
            description: Optional description.
            tags: Optional list of tags for filtering.
            uploaded_by: User ID of uploader.

        Returns:
            Dict with document info and chunk count.
        """
        # Detect file type
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if ext not in PARSERS:
            raise ValueError(f"Unsupported file type: .{ext}. Supported: {', '.join(PARSERS.keys())}")

        # Create document record
        doc = KBDocument(
            tenant_id=self.tenant_id,
            uploaded_by=uploaded_by,
            title=title or filename,
            filename=filename,
            file_type=ext,
            file_size=len(file_bytes),
            status="processing",
            category=category,
            description=description,
            tags=tags,
        )
        self.db.add(doc)
        await self.db.flush()  # Get the ID

        try:
            # Parse document — run in thread to avoid blocking the event loop
            # (PyPDF2 parsing is CPU-intensive and can trigger greenlet errors)
            parser = PARSERS[ext]
            content, page_count = await asyncio.to_thread(parser, file_bytes)
            doc.page_count = page_count

            # Strip null bytes — PostgreSQL rejects 0x00 in text columns
            content = content.replace("\x00", "")

            if not content.strip():
                doc.status = "failed"
                doc.error_message = "Document contains no extractable text"
                await self.db.commit()
                return {"id": doc.id, "status": "failed", "error_message": doc.error_message,
                        "filename": doc.filename, "title": doc.title, "chunk_count": 0}

            # Chunk document — also CPU-intensive for large docs
            chunk_size = int(os.getenv("RAG_CHUNK_SIZE", "1024"))
            chunk_overlap = int(os.getenv("RAG_CHUNK_OVERLAP", "200"))
            chunks = await asyncio.to_thread(chunk_text, content, chunk_size, chunk_overlap)

            if not chunks:
                doc.status = "failed"
                doc.error_message = "No chunks produced from document"
                await self.db.commit()
                return doc.to_dict()

            # Generate embeddings in batches (graceful degradation if service unavailable)
            batch_size = 32
            all_embeddings = []
            embedding_available = True
            try:
                for i in range(0, len(chunks), batch_size):
                    batch_texts = [c["content"] for c in chunks[i:i + batch_size]]
                    batch_embeddings = await self.embedding_service.embed_texts(batch_texts)
                    all_embeddings.extend(batch_embeddings)
            except (RuntimeError, Exception) as emb_err:
                logger.warning(
                    f"Embedding service unavailable — storing chunks without embeddings: {emb_err}"
                )
                embedding_available = False
                all_embeddings = []

            # Store chunks (with or without embeddings)
            for idx, chunk_data in enumerate(chunks):
                embedding = all_embeddings[idx] if embedding_available and idx < len(all_embeddings) else None
                chunk = KBChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_data["content"],
                    token_count=len(chunk_data["content"].split()),  # Approximate
                    start_char=chunk_data["start_char"],
                    end_char=chunk_data["end_char"],
                    embedding=embedding,
                )
                self.db.add(chunk)

            # Update document status
            doc.status = "indexed" if embedding_available else "pending_embedding"
            doc.chunk_count = len(chunks)
            if embedding_available and all_embeddings:
                doc.embedding_model = self.embedding_service.model
                doc.embedding_dimensions = len(all_embeddings[0])

            # Build result dict BEFORE commit to avoid post-commit greenlet issues
            # (server_default columns like created_at aren't loaded until refresh,
            # and refresh can trigger greenlet errors with mixed sync/async sessions)
            result = {
                "id": doc.id,
                "tenant_id": doc.tenant_id,
                "uploaded_by": doc.uploaded_by,
                "title": doc.title,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "page_count": doc.page_count,
                "status": doc.status,
                "error_message": doc.error_message,
                "chunk_count": doc.chunk_count,
                "embedding_model": doc.embedding_model,
                "embedding_dimensions": doc.embedding_dimensions,
                "category": doc.category,
                "description": doc.description,
                "tags": doc.tags,
                "created_at": None,
                "updated_at": None,
            }

            await self.db.commit()
            logger.info(
                f"Ingested document '{doc.title}' ({doc.file_type}): "
                f"{doc.chunk_count} chunks, {doc.page_count} pages"
            )
            return result

        except Exception as e:
            logger.error(f"Document ingestion failed for '{filename}': {e}")
            try:
                doc.status = "failed"
                doc.error_message = str(e)[:500]
                await self.db.commit()
            except Exception:
                pass  # Best-effort status update
            raise

    # ------------------------------------------------------------------
    # Programmatic Text Ingestion (for auto-generated content)
    # ------------------------------------------------------------------

    async def ingest_text(
        self,
        text_content: str,
        title: str,
        category: str = "general",
        description: Optional[str] = None,
        tags: Optional[List[str]] = None,
        uploaded_by: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Ingest plain text directly (no file parsing step).

        Convenience wrapper for programmatically generated content such as
        auto-indexed supply chain config descriptions.

        Args:
            text_content: The text to chunk and embed.
            title: Document title.
            category: Document category for filtering.
            description: Optional description.
            tags: Optional tags for filtering.
            uploaded_by: User ID of creator.

        Returns:
            Dict with document info and chunk count.
        """
        file_bytes = text_content.encode("utf-8")
        filename = f"{category}_{title[:60].replace(' ', '_')}.txt"
        return await self.ingest_document(
            file_bytes=file_bytes,
            filename=filename,
            title=title,
            category=category,
            description=description,
            tags=tags,
            uploaded_by=uploaded_by,
        )

    # ------------------------------------------------------------------
    # Bulk Deletion (for re-indexing)
    # ------------------------------------------------------------------

    async def delete_by_category_and_tag(self, category: str, tag: str) -> int:
        """Delete all documents matching a category and containing a specific tag.

        Used for re-indexing: delete old config docs before creating new ones.
        Cascading delete removes associated chunks automatically.

        Args:
            category: Document category (e.g., "supply_chain_config").
            tag: Tag string to match (e.g., "config_id:5").

        Returns:
            Number of documents deleted.
        """
        # Find matching documents
        stmt = select(KBDocument).where(
            KBDocument.tenant_id == self.tenant_id,
            KBDocument.category == category,
        )
        result = await self.db.execute(stmt)
        docs = result.scalars().all()

        deleted = 0
        for doc in docs:
            doc_tags = doc.tags or []
            if tag in doc_tags:
                await self.db.delete(doc)
                deleted += 1

        if deleted:
            await self.db.commit()
            logger.info(f"Deleted {deleted} KB doc(s) matching category={category}, tag={tag}")

        return deleted

    # ------------------------------------------------------------------
    # Document Management
    # ------------------------------------------------------------------

    async def list_documents(
        self,
        status: Optional[str] = None,
        category: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """List all documents for the customer."""
        stmt = select(KBDocument).where(KBDocument.tenant_id == self.tenant_id)
        if status:
            stmt = stmt.where(KBDocument.status == status)
        if category:
            stmt = stmt.where(KBDocument.category == category)
        stmt = stmt.order_by(KBDocument.created_at.desc())

        result = await self.db.execute(stmt)
        docs = result.scalars().all()
        return [d.to_dict() for d in docs]

    async def get_document(self, document_id: int) -> Optional[Dict[str, Any]]:
        """Get a single document by ID."""
        stmt = select(KBDocument).where(
            KBDocument.id == document_id,
            KBDocument.tenant_id == self.tenant_id,
        )
        result = await self.db.execute(stmt)
        doc = result.scalar_one_or_none()
        return doc.to_dict() if doc else None

    async def delete_document(self, document_id: int) -> bool:
        """Delete a document and all its chunks."""
        stmt = select(KBDocument).where(
            KBDocument.id == document_id,
            KBDocument.tenant_id == self.tenant_id,
        )
        result = await self.db.execute(stmt)
        doc = result.scalar_one_or_none()

        if not doc:
            return False

        await self.db.delete(doc)
        await self.db.commit()
        logger.info(f"Deleted document '{doc.title}' (id={doc.id})")
        return True

    # ------------------------------------------------------------------
    # RAG Search
    # ------------------------------------------------------------------

    async def search(
        self,
        query: str,
        top_k: Optional[int] = None,
        category: Optional[str] = None,
    ) -> List[ChunkResult]:
        """Semantic search across the knowledge base.

        Embeds the query, then finds the most similar document chunks
        using pgvector cosine distance.

        Args:
            query: Natural language search query.
            top_k: Number of results to return (default from config).
            category: Optional filter by document category.

        Returns:
            List of ChunkResult sorted by relevance (highest score first).
        """
        if top_k is None:
            top_k = int(os.getenv("RAG_TOP_K", "5"))

        # Embed the query
        query_embedding = await self.embedding_service.embed_query(query)

        # Build pgvector cosine similarity query
        # cosine_distance returns distance (0 = identical), we convert to similarity
        embedding_str = f"[{','.join(str(x) for x in query_embedding)}]"

        # Use raw SQL for pgvector operator — SQLAlchemy ORM doesn't natively
        # support the <=> operator without custom types
        sql = text("""
            SELECT
                c.id AS chunk_id,
                c.document_id,
                d.title AS document_title,
                c.content,
                c.chunk_index,
                c.page_number,
                c.metadata,
                1 - (c.embedding <=> CAST(:embedding AS vector)) AS score
            FROM kb_chunks c
            JOIN kb_documents d ON c.document_id = d.id
            WHERE d.tenant_id = :tenant_id
              AND d.status = 'indexed'
              {category_filter}
            ORDER BY c.embedding <=> CAST(:embedding AS vector)
            LIMIT :top_k
        """.format(
            category_filter="AND d.category = :category" if category else ""
        ))

        params: Dict[str, Any] = {
            "embedding": embedding_str,
            "tenant_id": self.tenant_id,
            "top_k": top_k,
        }
        if category:
            params["category"] = category

        result = await self.db.execute(sql, params)
        rows = result.fetchall()

        return [
            ChunkResult(
                chunk_id=row.chunk_id,
                document_id=row.document_id,
                document_title=row.document_title,
                content=row.content,
                chunk_index=row.chunk_index,
                page_number=row.page_number,
                score=float(row.score),
                metadata=row.metadata,
            )
            for row in rows
        ]

    async def search_for_context(
        self,
        query: str,
        top_k: int = 5,
        max_tokens: int = 4000,
    ) -> str:
        """Search and format results as context string for LLM injection.

        Returns a formatted string suitable for prepending to an LLM prompt.
        Truncates to max_tokens approximate word count.
        """
        results = await self.search(query, top_k=top_k)

        if not results:
            return ""

        context_parts = []
        total_words = 0
        for r in results:
            words = r.content.split()
            if total_words + len(words) > max_tokens:
                break
            source = f"[{r.document_title}"
            if r.page_number:
                source += f", p.{r.page_number}"
            source += f", relevance: {r.score:.2f}]"
            context_parts.append(f"{source}\n{r.content}")
            total_words += len(words)

        if not context_parts:
            return ""

        return (
            "=== KNOWLEDGE BASE CONTEXT ===\n"
            + "\n\n---\n\n".join(context_parts)
            + "\n=== END CONTEXT ===\n"
        )

    # ------------------------------------------------------------------
    # URL Ingestion
    # ------------------------------------------------------------------

    async def ingest_from_url(
        self,
        url: str,
        title: Optional[str] = None,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None,
        uploaded_by: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Fetch a URL and ingest its content into the knowledge base.

        Handles HTML pages (extracts readable text) and direct PDF/DOCX links.

        Args:
            url: HTTP/HTTPS URL to fetch.
            title: Human-readable title (defaults to page title or URL).
            category: Document category for filtering.
            tags: Optional list of tags.
            uploaded_by: User ID of requester.

        Returns:
            Dict with document info and chunk count.
        """
        import httpx
        from urllib.parse import urlparse

        try:
            async with httpx.AsyncClient(
                follow_redirects=True,
                timeout=30.0,
                headers={"User-Agent": "Mozilla/5.0 (compatible; Autonomy-RAG/1.0)"},
            ) as client:
                response = await client.get(url)
                response.raise_for_status()
        except httpx.HTTPStatusError as e:
            raise ValueError(f"HTTP {e.response.status_code} fetching URL: {url}")
        except httpx.RequestError as e:
            raise ValueError(f"Failed to fetch URL: {e}")

        content_type = response.headers.get("content-type", "").lower()
        raw_bytes = response.content
        parsed = urlparse(url)
        url_filename = parsed.path.split("/")[-1] or "page.html"

        if "pdf" in content_type or url_filename.lower().endswith(".pdf"):
            filename = url_filename if url_filename.lower().endswith(".pdf") else "document.pdf"
        elif "docx" in content_type or url_filename.lower().endswith(".docx"):
            filename = url_filename
        elif "text/plain" in content_type or url_filename.lower().endswith(".txt"):
            filename = url_filename
        else:
            # Treat as HTML — extract readable text
            html_text = _extract_html_text(response.text)
            if not html_text.strip():
                raise ValueError("Page returned no readable text content")
            if not title:
                title = _extract_html_title(response.text) or url
            raw_bytes = html_text.encode("utf-8")
            filename = "page.txt"

        effective_title = title or url_filename or url
        if not tags:
            tags = []
        if "url_source" not in tags:
            tags = ["url_source"] + tags

        return await self.ingest_document(
            file_bytes=raw_bytes,
            filename=filename,
            title=effective_title,
            category=category,
            tags=tags,
            uploaded_by=uploaded_by,
        )

    # ------------------------------------------------------------------
    # Status / Statistics
    # ------------------------------------------------------------------

    async def get_status(self) -> Dict[str, Any]:
        """Get knowledge base statistics for the customer."""
        doc_count = await self.db.execute(
            select(func.count(KBDocument.id)).where(KBDocument.tenant_id == self.tenant_id)
        )
        chunk_count = await self.db.execute(
            select(func.count(KBChunk.id)).join(KBDocument).where(
                KBDocument.tenant_id == self.tenant_id
            )
        )
        indexed_count = await self.db.execute(
            select(func.count(KBDocument.id)).where(
                KBDocument.tenant_id == self.tenant_id,
                KBDocument.status == "indexed",
            )
        )

        embedding_health = await self.embedding_service.health_check()

        return {
            "total_documents": doc_count.scalar() or 0,
            "indexed_documents": indexed_count.scalar() or 0,
            "total_chunks": chunk_count.scalar() or 0,
            "embedding_service": embedding_health,
            "rag_enabled": os.getenv("RAG_ENABLED", "false").lower() == "true",
            "chunk_size": int(os.getenv("RAG_CHUNK_SIZE", "1024")),
            "chunk_overlap": int(os.getenv("RAG_CHUNK_OVERLAP", "200")),
            "top_k": int(os.getenv("RAG_TOP_K", "5")),
        }
